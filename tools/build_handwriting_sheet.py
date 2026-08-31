from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(__file__).resolve().parents[1]
OUT=ROOT/'templates'/'KC_DP2_Persoenliches_Handschriftprofil_Lernbogen_V1.docx'
BURGUNDY='741521'; GOLD='B58A42'; LIGHT='F7F3EE'; LINE='D8CFC7'; INK='292522'; MUTED='6D655F'

def shade(cell,color):
    tcPr=cell._tc.get_or_add_tcPr();shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd');tcPr.append(shd)
    shd.set(qn('w:fill'),color)

def margins(cell,top=70,start=90,bottom=70,end=90):
    tcPr=cell._tc.get_or_add_tcPr();tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar');tcPr.append(tcMar)
    for tag,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        el=tcMar.find(qn('w:'+tag))
        if el is None: el=OxmlElement('w:'+tag);tcMar.append(el)
        el.set(qn('w:w'),str(val));el.set(qn('w:type'),'dxa')

def set_cell_width(cell,dxa):
    tcPr=cell._tc.get_or_add_tcPr();tcW=tcPr.find(qn('w:tcW'))
    if tcW is None: tcW=OxmlElement('w:tcW');tcPr.append(tcW)
    tcW.set(qn('w:w'),str(dxa));tcW.set(qn('w:type'),'dxa')

def table_geometry(table,widths):
    table.autofit=False;table.alignment=WD_TABLE_ALIGNMENT.CENTER
    tblPr=table._tbl.tblPr;tblW=tblPr.find(qn('w:tblW'))
    if tblW is None: tblW=OxmlElement('w:tblW');tblPr.append(tblW)
    tblW.set(qn('w:w'),str(sum(widths)));tblW.set(qn('w:type'),'dxa')
    grid=table._tbl.tblGrid
    for c in list(grid): grid.remove(c)
    for width in widths:
        col=OxmlElement('w:gridCol');col.set(qn('w:w'),str(width));grid.append(col)
    for row in table.rows:
        for cell,width in zip(row.cells,widths): set_cell_width(cell,width);margins(cell)

def font(run,size=10,bold=False,color=INK,name='Aptos'):
    run.font.name=name;run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'),name);run._element.rPr.rFonts.set(qn('w:hAnsi'),name)
    run.font.size=Pt(size);run.bold=bold;run.font.color.rgb=RGBColor.from_string(color)

def style_doc(doc):
    sec=doc.sections[0];sec.page_width=Inches(8.5);sec.page_height=Inches(11);sec.top_margin=Inches(.55);sec.bottom_margin=Inches(.55);sec.left_margin=Inches(.55);sec.right_margin=Inches(.55);sec.header_distance=Inches(.25);sec.footer_distance=Inches(.28)
    normal=doc.styles['Normal'];normal.font.name='Aptos';normal._element.rPr.rFonts.set(qn('w:ascii'),'Aptos');normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Aptos');normal.font.size=Pt(10);normal.font.color.rgb=RGBColor.from_string(INK);normal.paragraph_format.space_after=Pt(4);normal.paragraph_format.line_spacing=1.05
    for name,size,before,after in [('Heading 1',16,9,5),('Heading 2',12,7,4)]:
        s=doc.styles[name];s.font.name='Aptos';s._element.rPr.rFonts.set(qn('w:ascii'),'Aptos');s._element.rPr.rFonts.set(qn('w:hAnsi'),'Aptos');s.font.size=Pt(size);s.font.bold=True;s.font.color.rgb=RGBColor.from_string(BURGUNDY);s.paragraph_format.space_before=Pt(before);s.paragraph_format.space_after=Pt(after)

def footer(section):
    p=section.footer.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run('KC DP2 · Persönliches Handschriftprofil · Lernbogen V1'),8,False,MUTED)

def title(doc,subtitle):
    t=doc.add_table(rows=1,cols=2);table_geometry(t,[1200,9300])
    c=t.cell(0,0);shade(c,LIGHT);c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    try:p.add_run().add_picture(str(ROOT/'assets'/'kc-dp2-icon-192.png'),width=Inches(.52))
    except Exception:font(p.add_run('KC'),18,True,BURGUNDY)
    c=t.cell(0,1);shade(c,LIGHT);p=c.paragraphs[0];font(p.add_run('KC DP2 · Persönliches Handschriftprofil'),18,True,BURGUNDY);p.paragraph_format.space_after=Pt(2)
    p=c.add_paragraph();font(p.add_run(subtitle),10,True,GOLD);p.paragraph_format.space_after=Pt(0)

def info_block(doc):
    t=doc.add_table(rows=2,cols=4);table_geometry(t,[1500,3750,1500,3750])
    data=[('Profil-ID','____________________________','Bogen-ID','wird beim Einlesen erzeugt'),('Datum','____________________________','Einwilligung','☐ bestätigt  ☐ noch offen')]
    for r,row in enumerate(data):
        for c,val in enumerate(row):
            cell=t.cell(r,c);shade(cell,LIGHT if c%2==0 else 'FFFFFF');p=cell.paragraphs[0];font(p.add_run(val),9,c%2==0,BURGUNDY if c%2==0 else INK)

def instruction(doc,text):
    t=doc.add_table(rows=1,cols=1);table_geometry(t,[10500]);shade(t.cell(0,0),'FFF8E9');p=t.cell(0,0).paragraphs[0];font(p.add_run('Hinweis: '),9,True,GOLD);font(p.add_run(text),9,False,INK)

def practice_grid(doc,labels,repeats=3,cols=5,box_height=390):
    row_count=(len(labels)+cols-1)//cols
    t=doc.add_table(rows=row_count,cols=cols);table_geometry(t,[10500//cols]*cols)
    for i,cell in enumerate(c for row in t.rows for c in row.cells):
        cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=cell.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(0)
        if i<len(labels):
            shade(cell,LIGHT);font(p.add_run(f'{labels[i]}  '),11,True,BURGUNDY)
            font(p.add_run('  '.join('________' for _ in range(repeats))),8,False,'A9A19A')
        else: shade(cell,'FFFFFF')
    return t

def phrase_grid(doc,phrases):
    t=doc.add_table(rows=len(phrases),cols=2);table_geometry(t,[3150,7350])
    for r,text in enumerate(phrases):
        shade(t.cell(r,0),LIGHT);p=t.cell(r,0).paragraphs[0];font(p.add_run(text),10,True,BURGUNDY)
        p=t.cell(r,1).paragraphs[0];p.paragraph_format.space_after=Pt(6);font(p.add_run('Bitte hier deutlich schreiben:'),7,False,'B8B1AA');p.add_run('\n');font(p.add_run('________________________________________________________'),9,False,'A9A19A')

def page_break(doc):
    doc.add_page_break()

doc=Document();style_doc(doc);footer(doc.sections[0])
title(doc,'Druck- und Lernbogen · keine automatische Übernahme ohne Prüfung')
info_block(doc);instruction(doc,'Bitte mit demselben Stift wie auf den Wunschformularen schreiben. Jedes Feld einzeln und vollständig ausfüllen. Der QR-/Profilbezug enthält keinen ausgeschriebenen Namen.')
doc.add_heading('1 · Ziffern und typische Verwechslungen',level=1)
practice_grid(doc,list('0123456789'),3)
doc.add_heading('Verwechslungsgruppen',level=2)
phrase_grid(doc,['0  O','1  I  l','2  Z','5  S','6  G','9  g'])

page_break(doc);title(doc,'Buchstaben und Bereichskürzel')
instruction(doc,'Groß- und Kleinbuchstaben jeweils in die freien Felder schreiben. Umlaute und ß bitte besonders deutlich ausführen.')
doc.add_heading('2 · Großbuchstaben A–Z',level=1);practice_grid(doc,list('ABCDEFGHIJKLMNOPQRSTUVWXYZ'),2,box_height=300)
doc.add_heading('3 · Kleinbuchstaben und Sonderzeichen',level=1);practice_grid(doc,list('abcdefghijklmnopqrstuvwxyz')+['Ä','Ö','Ü','ä','ö','ü','ß'],2,box_height=270)
doc.add_heading('Bereiche',level=2);phrase_grid(doc,['V = Vorne','H = Hinten','B = Beides','Z = Zusatzbereich'])

page_break(doc);title(doc,'Uhrzeiten, Datum und DP2-Praxisbegriffe')
instruction(doc,'Diese Beispiele helfen der Erkennung im praktischen Einsatz. Bitte exakt den links vorgegebenen Inhalt rechts abschreiben.')
doc.add_heading('4 · Uhrzeit und Datum',level=1);phrase_grid(doc,['08:00','14:30','18:00–23:00','04.12.2026'])
doc.add_heading('5 · Status und Planung',level=1);phrase_grid(doc,['frei / Urlaub','gesperrt / Reserve','nur wenn nötig','Aufbau / Nachbereitung','Wunsch 14:00–16:00','Kann 14:00–20:00'])
doc.add_heading('6 · Qualitätskontrolle',level=1)
t=doc.add_table(rows=4,cols=2);table_geometry(t,[3600,6900])
for r,(a,b) in enumerate([('Vollständigkeit','☐ vollständig  ☐ Nachlernen nötig'),('Stift/Lesbarkeit','☐ gut  ☐ mittel  ☐ unzureichend'),('Kontrolliert durch','____________________________________________'),('Freigabe Profil','☐ noch nicht  ☐ einsatzbereit')]):
    shade(t.cell(r,0),LIGHT);font(t.cell(r,0).paragraphs[0].add_run(a),9,True,BURGUNDY);font(t.cell(r,1).paragraphs[0].add_run(b),9,False,INK)
instruction(doc,'Erkannte Schriftproben werden erst nach ausdrücklicher Bestätigung gespeichert. Das Profil kann gesperrt, nachgelernt oder vollständig gelöscht werden.')

OUT.parent.mkdir(parents=True,exist_ok=True);doc.save(OUT);print(OUT)
